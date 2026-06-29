from docx import Document
from docx.shared import Pt


OUTPUT = "content/lesson_06.docx"


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    return table


def add_label_paragraph(document, label, text=""):
    paragraph = document.add_paragraph()
    paragraph.add_run(label).bold = True
    if text:
        paragraph.add_run(text)
    return paragraph


def add_code(document, code):
    for line in code.strip("\n").splitlines():
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10)


def add_numbered(document, items):
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


doc = Document()
doc.add_heading("Урок 6 — Наследование и переопределение: общий родитель и разное поведение", level=1)

doc.add_heading("Общая информация", level=2)
add_table(
    doc,
    ["Параметр", "Значение"],
    [
        ["Курс", "Программирование на C#"],
        ["Модуль", "Введение в ООП"],
        ["Номер урока", "6"],
        ["Возраст учащихся", "12-15 лет"],
        ["Продолжительность", "120 мин"],
    ],
)

doc.add_heading("Цель урока", level=2)
doc.add_paragraph(
    "К концу урока ученики смогут создать базовый класс `Enemy`, два класса-наследника, передать данные в родителя через `base` "
    "и переопределить метод `Attack()` через `virtual` и `override` так, чтобы разные объекты выполняли одну команду по-разному."
)

doc.add_heading("План урока", level=2)
add_table(
    doc,
    ["Этап", "Время"],
    [
        ["1. Организационный момент", "5 мин"],
        ["2. Теоретическая часть", "10 мин"],
        ["3. Практическая работа", "65 мин"],
        ["4. Самостоятельная работа", "30 мин"],
        ["5. Подведение итогов", "10 мин"],
        ["Итого", "120 мин"],
    ],
)

doc.add_heading("Ход занятия", level=2)

doc.add_heading("1. Организационный момент (5 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя:")
add_bullets(
    doc,
    [
        "Проверить, что у учеников открыт Visual Studio 2022 и доступен проект прошлого урока или новый проект `OopLesson6`.",
        "Попросить открыть `Program.cs` через `Solution Explorer`. Если окно закрыто, открыть его через `View` -> `Solution Explorer`.",
        "Повторить урок 5: несколько объектов можно хранить в `List<Enemy>` или `List<Hero>` и проходить по ним через `foreach`.",
        "Задать вопрос урока: что делать, если похожие объекты имеют общие данные, но некоторые действия выполняют по-разному?",
        "Назвать результат: сегодня ученики сделают общего родителя `Enemy`, несколько наследников и разные версии метода `Attack()`.",
    ],
)

doc.add_heading("2. Теоретическая часть (10 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя:")
doc.add_paragraph(
    "Это урок в режиме B: не давать длинную теорию заранее. В начале только ввести идею \"общий родитель - разные потомки\", "
    "а синтаксис `:`, `base`, `virtual` и `override` раскрывать прямо в практических шагах. На доске достаточно схемы: `Enemy` сверху, "
    "ниже `FastEnemy` и `HeavyEnemy`."
)

doc.add_heading("Общий родитель и потомки", level=4)
doc.add_paragraph(
    "Если несколько классов очень похожи, одинаковые свойства и методы можно вынести в общий базовый класс. "
    "Например, разные противники могут иметь `Name`, `Health`, `Damage` и общий метод `ShowInfo()`."
)
doc.add_paragraph(
    "Дочерний класс получает общее от родителя и может добавить свое. Простая формулировка: родитель хранит общее, потомок уточняет, чем он отличается."
)

doc.add_heading("Одна команда, разное поведение", level=4)
doc.add_paragraph(
    "Иногда у разных объектов есть действие с одинаковым смыслом, но результат должен отличаться. Все противники могут атаковать, "
    "но быстрый противник и тяжелый противник делают это по-разному. В C# для этого используют `virtual` в родителе и `override` в потомке."
)

doc.add_heading("Записи в блокнот", level=4)
for term, definition in [
    ("Наследование", "способ создать новый класс на основе уже существующего класса."),
    ("Базовый класс", "общий родитель, где лежат одинаковые данные и методы."),
    ("Дочерний класс", "класс-потомок, который получает возможности базового класса."),
    ("`:`", "знак, который показывает наследование: `class FastEnemy : Enemy`."),
    ("`base`", "обращение к родительскому конструктору или родительскому поведению."),
    ("`virtual` и `override`", "`virtual` разрешает заменить метод, `override` делает свою версию метода."),
]:
    paragraph = doc.add_paragraph()
    paragraph.add_run(term).bold = True
    paragraph.add_run(" — " + definition)

doc.add_heading("3. Практическая работа (65 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя:")
doc.add_paragraph(
    "Практику строить как постепенное усложнение одного проекта. Сначала ученики видят повторение кода в двух похожих классах. "
    "Затем преподаватель показывает базовый класс `Enemy`, группа вместе добавляет наследников, после этого метод `Attack()` становится `virtual`, "
    "а наследники получают свои версии через `override`. В конце список `List<Enemy>` показывает главную идею: одна команда вызывает разное поведение."
)

doc.add_heading("Задание", level=4)
doc.add_paragraph("Мини-теория 1. Зачем нужен базовый класс.")
doc.add_paragraph(
    "Если два класса содержат одинаковые свойства `Name`, `Health`, `Damage`, одинаковый конструктор и одинаковый `ShowInfo()`, код начинает повторяться. "
    "Повторение опасно: если нужно изменить вывод здоровья, придется искать одинаковые места в нескольких классах."
)
add_numbered(
    doc,
    [
        "Открыть Visual Studio 2022.",
        "Создать новый проект `Console App` с именем `OopLesson6` или открыть учебный проект.",
        "Открыть файл `Program.cs`.",
        "Вставить стартовый код с двумя похожими классами и запустить программу.",
        "Найти глазами повторяющиеся свойства и метод `ShowInfo()`.",
    ],
)
doc.add_paragraph("[СКРИНШОТ: Visual Studio 2022 — Program.cs с двумя похожими классами FastEnemy и HeavyEnemy до наследования]")
doc.add_paragraph("Стартовый код для обсуждения повторения:")
add_code(
    doc,
    '''
FastEnemy fastEnemy = new FastEnemy("Быстрый противник", 50, 8);
HeavyEnemy heavyEnemy = new HeavyEnemy("Тяжелый противник", 120, 18);

fastEnemy.ShowInfo();
heavyEnemy.ShowInfo();

class FastEnemy
{
    public string Name { get; private set; }
    public int Health { get; private set; }
    public int Damage { get; private set; }

    public FastEnemy(string name, int health, int damage)
    {
        Name = name;
        Health = health;
        Damage = damage;
    }

    public void ShowInfo()
    {
        Console.WriteLine(Name + ": здоровье " + Health + ", урон " + Damage);
    }
}

class HeavyEnemy
{
    public string Name { get; private set; }
    public int Health { get; private set; }
    public int Damage { get; private set; }

    public HeavyEnemy(string name, int health, int damage)
    {
        Name = name;
        Health = health;
        Damage = damage;
    }

    public void ShowInfo()
    {
        Console.WriteLine(Name + ": здоровье " + Health + ", урон " + Damage);
    }
}
''',
)

doc.add_paragraph("Мини-теория 2. Создаем общего родителя `Enemy`.")
doc.add_paragraph(
    "Общее переносим в `Enemy`: свойства, конструктор и метод `ShowInfo()`. Дочерний класс записывается через двоеточие: "
    "`class FastEnemy : Enemy`. В конструкторе дочернего класса строка `: base(name, health, damage)` передает данные в конструктор родителя."
)
add_numbered(
    doc,
    [
        "Удалить дублирующиеся свойства из `FastEnemy` и `HeavyEnemy`.",
        "Создать общий класс `Enemy`.",
        "Перенести в `Enemy` свойства `Name`, `Health`, `Damage`.",
        "Перенести в `Enemy` конструктор и `ShowInfo()`.",
        "Сделать `FastEnemy : Enemy` и `HeavyEnemy : Enemy`.",
        "В конструкторах наследников вызвать `base(name, health, damage)`.",
        "Запустить программу и убедиться, что `ShowInfo()` работает у объектов-наследников.",
    ],
)
doc.add_paragraph("[СКРИНШОТ: Visual Studio 2022 — выделены строки class FastEnemy : Enemy и : base(name, health, damage)]")
doc.add_paragraph("Код после выделения общего родителя:")
add_code(
    doc,
    '''
FastEnemy fastEnemy = new FastEnemy("Быстрый противник", 50, 8);
HeavyEnemy heavyEnemy = new HeavyEnemy("Тяжелый противник", 120, 18);

fastEnemy.ShowInfo();
heavyEnemy.ShowInfo();

class Enemy
{
    public string Name { get; private set; }
    public int Health { get; private set; }
    public int Damage { get; private set; }

    public Enemy(string name, int health, int damage)
    {
        Name = name;
        Health = health;
        Damage = damage;
    }

    public void ShowInfo()
    {
        Console.WriteLine(Name + ": здоровье " + Health + ", урон " + Damage);
    }
}

class FastEnemy : Enemy
{
    public FastEnemy(string name, int health, int damage)
        : base(name, health, damage)
    {
    }
}

class HeavyEnemy : Enemy
{
    public HeavyEnemy(string name, int health, int damage)
        : base(name, health, damage)
    {
    }
}
''',
)
doc.add_paragraph("[СКРИНШОТ: консоль — данные двух объектов выводятся через общий метод ShowInfo() из Enemy]")

doc.add_paragraph("Мини-теория 3. `virtual` и `override`.")
doc.add_paragraph(
    "`virtual` пишется у метода в родителе и означает: этот метод можно заменить в дочернем классе. "
    "`override` пишется у метода в потомке и означает: здесь своя версия метода родителя. Не вводить сложную теорию полиморфизма; достаточно фразы "
    "\"одна команда - разное поведение\"."
)
add_numbered(
    doc,
    [
        "В классе `Enemy` добавить метод `Attack()`.",
        "Сделать метод родителя виртуальным: `public virtual void Attack()`.",
        "В `FastEnemy` написать `public override void Attack()`.",
        "В `HeavyEnemy` написать `public override void Attack()`.",
        "В каждой версии вывести разный текст атаки.",
        "Создать два объекта и вызвать `Attack()` у каждого.",
        "Запустить программу и сравнить строки в консоли.",
    ],
)
doc.add_paragraph("[СКРИНШОТ: Visual Studio 2022 — выделены слова virtual в Enemy и override в FastEnemy и HeavyEnemy]")
doc.add_paragraph("Код с переопределением:")
add_code(
    doc,
    '''
FastEnemy fastEnemy = new FastEnemy("Быстрый противник", 50, 8);
HeavyEnemy heavyEnemy = new HeavyEnemy("Тяжелый противник", 120, 18);

fastEnemy.Attack();
heavyEnemy.Attack();

class Enemy
{
    public string Name { get; private set; }
    public int Health { get; private set; }
    public int Damage { get; private set; }

    public Enemy(string name, int health, int damage)
    {
        Name = name;
        Health = health;
        Damage = damage;
    }

    public void ShowInfo()
    {
        Console.WriteLine(Name + ": здоровье " + Health + ", урон " + Damage);
    }

    public virtual void Attack()
    {
        Console.WriteLine(Name + " атакует и наносит " + Damage + " урона.");
    }
}

class FastEnemy : Enemy
{
    public FastEnemy(string name, int health, int damage)
        : base(name, health, damage)
    {
    }

    public override void Attack()
    {
        Console.WriteLine(Name + " быстро атакует и наносит " + Damage + " урона.");
    }
}

class HeavyEnemy : Enemy
{
    public HeavyEnemy(string name, int health, int damage)
        : base(name, health, damage)
    {
    }

    public override void Attack()
    {
        Console.WriteLine(Name + " наносит тяжелый удар на " + Damage + " урона.");
    }
}
''',
)
doc.add_paragraph("[СКРИНШОТ: консоль — два объекта вызывают Attack(), но выводят разные варианты атаки]")

doc.add_paragraph("Мини-теория 4. Один список, разные объекты.")
doc.add_paragraph(
    "Объекты-наследники можно положить в список базового типа: `List<Enemy>`. Для учеников это важная практическая картина: "
    "в одном списке лежат разные виды противников, но у всех есть общий родитель `Enemy` и общая команда `Attack()`."
)
add_numbered(
    doc,
    [
        "Добавить в начало файла `using System.Collections.Generic;`.",
        "Создать список `List<Enemy> enemies = new List<Enemy>();`.",
        "Добавить в список `new FastEnemy(...)` и `new HeavyEnemy(...)`.",
        "Написать `foreach (Enemy enemy in enemies)`.",
        "Внутри цикла вызвать `enemy.ShowInfo();` и `enemy.Attack();`.",
        "Запустить программу и увидеть, что одна команда `enemy.Attack()` дает разный результат.",
    ],
)
doc.add_paragraph("[СКРИНШОТ: Visual Studio 2022 — List<Enemy> enemies, Add() с разными наследниками и foreach с enemy.Attack()]")
doc.add_paragraph("Финальный код практики:")
add_code(
    doc,
    '''
using System.Collections.Generic;

List<Enemy> enemies = new List<Enemy>();

enemies.Add(new FastEnemy("Быстрый противник", 50, 8));
enemies.Add(new HeavyEnemy("Тяжелый противник", 120, 18));

foreach (Enemy enemy in enemies)
{
    enemy.ShowInfo();
    enemy.Attack();
    Console.WriteLine();
}

class Enemy
{
    public string Name { get; private set; }
    public int Health { get; private set; }
    public int Damage { get; private set; }

    public Enemy(string name, int health, int damage)
    {
        Name = name;
        Health = health;
        Damage = damage;
    }

    public void ShowInfo()
    {
        Console.WriteLine(Name + ": здоровье " + Health + ", урон " + Damage);
    }

    public virtual void Attack()
    {
        Console.WriteLine(Name + " атакует и наносит " + Damage + " урона.");
    }
}

class FastEnemy : Enemy
{
    public FastEnemy(string name, int health, int damage)
        : base(name, health, damage)
    {
    }

    public override void Attack()
    {
        Console.WriteLine(Name + " быстро атакует и наносит " + Damage + " урона.");
    }
}

class HeavyEnemy : Enemy
{
    public HeavyEnemy(string name, int health, int damage)
        : base(name, health, damage)
    {
    }

    public override void Attack()
    {
        Console.WriteLine(Name + " наносит тяжелый удар на " + Damage + " урона.");
    }
}
''',
)
doc.add_paragraph("[СКРИНШОТ: консоль — foreach вызывает ShowInfo() и разные версии Attack() для объектов из одного List<Enemy>]")
add_label_paragraph(
    doc,
    "Ожидаемый результат:",
    " у ученика есть проект, где `Enemy` хранит общее, наследники используют `base`, метод `Attack()` переопределен через `override`, а `List<Enemy>` запускает разные версии атаки через один `foreach`.",
)

doc.add_heading("4. Самостоятельная работа (30 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя:")
doc.add_paragraph(
    "Дать ученикам самостоятельное добавление третьего наследника. Первые 5 минут не диктовать код полностью: ученики должны попробовать повторить структуру "
    "`class NewEnemy : Enemy` и `: base(...)`. Затем пройти по классу и проверить три точки: наследование через двоеточие, вызов `base`, `override Attack()`."
)

doc.add_heading("Задание", level=4)
doc.add_paragraph("Добавь третий класс-наследник для `Enemy` и сделай свою версию атаки.")
add_numbered(
    doc,
    [
        "Придумай название класса: например `FireEnemy`, `IceEnemy`, `ShadowEnemy`, `RobotEnemy` или свой вариант.",
        "Создай класс-наследник: `class FireEnemy : Enemy`.",
        "Добавь конструктор с параметрами `name`, `health`, `damage`.",
        "Передай параметры родителю через `: base(name, health, damage)`.",
        "Переопредели метод `Attack()` через `public override void Attack()`.",
        "Внутри `Attack()` выведи свой текст атаки.",
        "Добавь новый объект в `List<Enemy>` через `Add()`.",
        "Запусти программу и проверь, что новый объект выводится в `foreach` и атакует по-своему.",
    ],
)
doc.add_paragraph("Каркас для самостоятельной работы:")
add_code(
    doc,
    '''
class FireEnemy : Enemy
{
    public FireEnemy(string name, int health, int damage)
        : base(name, health, damage)
    {
    }

    public override void Attack()
    {
        Console.WriteLine(Name + " атакует огнем и наносит " + Damage + " урона.");
    }
}
''',
)
doc.add_paragraph("Строка для добавления в список:")
add_code(
    doc,
    '''
enemies.Add(new FireEnemy("Огненный противник", 70, 14));
''',
)

doc.add_heading("Критерии оценки", level=4)
add_table(
    doc,
    ["Результат", "Оценка"],
    [
        ["Создан новый наследник `Enemy`, есть корректный конструктор с `base`, метод `Attack()` переопределен через `override`, объект добавлен в `List<Enemy>` и программа запускается", "Отлично"],
        ["Наследник и `override Attack()` работают, но есть небольшие ошибки в тексте вывода, названиях или оформлении кода", "Хорошо"],
        ["Наследник создан частично; код заработал после подсказок по `base`, скобкам или `override`", "Удовлетворительно"],
        ["Нет рабочего наследника или программа не запускается даже после базовой помощи", "Требует доработки"],
    ],
)

doc.add_heading("5. Подведение итогов (10 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя:")
add_bullets(
    doc,
    [
        "Попросить нескольких учеников показать свой третий класс-наследник и объяснить, что он получает от `Enemy`.",
        "Вернуться к формуле урока: родитель хранит общее, потомки делают свои версии поведения.",
        "Сравнить `virtual` и `override`: где пишется каждое слово и зачем.",
        "Показать ключевую строку `foreach (Enemy enemy in enemies)` и спросить, почему внутри списка могут лежать разные наследники.",
        "Анонсировать следующий урок: будем собирать отдельные классы в мини-проект с героем, противниками и предметами.",
    ],
)
add_label_paragraph(doc, "Вопросы для рефлексии:")
add_bullets(
    doc,
    [
        "Что лучше хранить в базовом классе `Enemy`: общее или уникальное?",
        "Что означает запись `class FastEnemy : Enemy`?",
        "Зачем в конструкторе наследника нужен `base(...)`?",
        "Где пишется `virtual`: в родителе или в потомке?",
        "Где пишется `override`: в родителе или в потомке?",
        "Почему один вызов `enemy.Attack()` может дать разный результат?",
    ],
)

doc.add_heading("Домашнее задание", level=2)
doc.add_paragraph("Вариант без компьютера:")
add_numbered(
    doc,
    [
        "Нарисовать схему классов: сверху общий `Enemy`, ниже три наследника.",
        "Подписать, какие свойства и методы лежат в `Enemy`.",
        "Подписать, чем отличается каждый наследник.",
        "Для каждого наследника написать одну фразу, как он выполняет `Attack()`.",
        "Отметить, где в схеме родитель, а где дочерние классы.",
    ],
)
doc.add_paragraph("Вариант с компьютером:")
add_numbered(
    doc,
    [
        "Открыть проект `OopLesson6`.",
        "Добавить четвертый класс-наследник от `Enemy`.",
        "Сделать для него собственный `override Attack()`.",
        "Добавить объект нового класса в `List<Enemy>`.",
        "Запустить программу и проверить вывод в консоли.",
    ],
)

doc.add_heading("Методические заметки преподавателя", level=2)
doc.add_heading("Возможные сложности", level=3)
add_bullets(
    doc,
    [
        "Ученики могут путать базовый и дочерний класс. Возвращайте к схеме: общее сверху, конкретные варианты ниже.",
        "В строке `class FastEnemy : Enemy` часто забывают двоеточие или пишут классы в неправильном порядке.",
        "Вызов `base(name, health, damage)` часто ломается из-за разного количества параметров в дочернем конструкторе и конструкторе `Enemy`.",
        "Ученики могут написать `override`, но забыть `virtual` в родительском методе. Visual Studio покажет ошибку; используйте ее как подсказку.",
        "Иногда ученики меняют название метода: в родителе `Attack`, а в потомке `Attak` или `FireAttack`. Подчеркните, что переопределение требует одинакового имени метода.",
        "В `List<Enemy>` ученики могут попытаться указать список конкретного наследника и удивиться, почему другой наследник туда не добавляется. Для общего списка нужен тип родителя.",
        "Слабые навыки Windows проявятся в том, что ученик откроет не тот проект или отдельный `Program.cs`. Проверяйте `Solution Explorer` и имя проекта `OopLesson6`.",
    ],
)

doc.add_heading("Способы помощи", level=3)
add_bullets(
    doc,
    [
        "Если ученик не понимает наследование, спросите: что одинакового у всех противников? Это и должно быть в родителе.",
        "Если ученик застрял на `base`, дайте подсказку: дочерний класс передает имя, здоровье и урон в родительский конструктор.",
        "Если не работает `override`, попросите сравнить строку метода в родителе и потомке: имя, скобки, параметры и слово `virtual`.",
        "Если в классе много ошибок, сначала проверьте фигурные скобки классов, затем первую ошибку Visual Studio сверху.",
        "Если ученик не понимает `List<Enemy>`, спросите: какой общий тип есть у всех объектов в списке?",
        "Если ученик просит готовый ответ, дайте каркас класса с пустым названием и пустым текстом атаки, чтобы он сам заполнил отличающуюся часть.",
        "Если группа быстро уходит в сложную теорию, мягко верните фокус: на этом уроке нужны только `:`, `base`, `virtual`, `override` и простой `foreach`.",
    ],
)

doc.add_heading("Дополнительные задания для тех, кто справился раньше", level=3)
add_bullets(
    doc,
    [
        "Добавить наследнику новое свойство, например `Speed`, `Armor` или `Element`, и вывести его в отдельном методе.",
        "Создать метод `Taunt()` в `Enemy`, сделать его `virtual` и переопределить в одном наследнике.",
        "Добавить в `List<Enemy>` четыре объекта и посчитать, сколько из них имеют `Damage` больше 10.",
        "Сделать метод `ShowEnemyTeam(List<Enemy> enemies)`, который выводит всех противников и вызывает их `Attack()`.",
        "Написать короткое объяснение для соседа: почему `List<Enemy>` может хранить объекты классов-наследников.",
    ],
)

doc.save(OUTPUT)
print(OUTPUT)
