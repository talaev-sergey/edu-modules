from docx import Document
from docx.shared import Pt


OUTPUT = "content/lesson_04.docx"


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    return table


def add_label_paragraph(document, label, text=""):
    paragraph = document.add_paragraph()
    paragraph.add_run(label).bold = True
    if text:
        paragraph.add_run(text)
    return paragraph


def add_code(document, code):
    for line in code.strip("\n").splitlines():
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10)


def add_numbered(document, items):
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


doc = Document()
doc.add_heading("Урок 4 — Инкапсуляция и свойства: защищаем данные объекта", level=1)

doc.add_heading("Общая информация", level=2)
add_table(
    doc,
    ["Параметр", "Значение"],
    [
        ["Курс", "Программирование на C#"],
        ["Модуль", "Введение в ООП"],
        ["Номер урока", "4"],
        ["Возраст учащихся", "12-15 лет"],
        ["Продолжительность", "120 мин"],
    ],
)

doc.add_heading("Цель урока", level=2)
doc.add_paragraph(
    "К концу урока ученики смогут закрыть внутренние данные объекта через `private`, изменить их через безопасные методы "
    "или свойства и добавить простую проверку значения."
)

doc.add_heading("План урока", level=2)
add_table(
    doc,
    ["Этап", "Время"],
    [
        ["1. Организационный момент", "5 мин"],
        ["2. Теоретическая часть", "25 мин"],
        ["3. Практическая работа", "50 мин"],
        ["4. Самостоятельная работа", "30 мин"],
        ["5. Подведение итогов", "10 мин"],
        ["Итого", "120 мин"],
    ],
)

doc.add_heading("Ход занятия", level=2)

doc.add_heading("1. Организационный момент (5 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя:")
add_bullets(
    doc,
    [
        "Проверить, что у учеников открыт Visual Studio 2022 и доступен проект прошлого урока или новый проект `OopLesson4`.",
        "Попросить открыть `Program.cs` через `Solution Explorer`. Если окна нет, открыть его через `View` -> `Solution Explorer`.",
        "Коротко повторить: конструктор задает начальное состояние объекта, поля хранят данные, методы меняют состояние.",
        "Показать вопрос урока: что будет, если любой код сможет напрямую поставить здоровью героя значение `-999`?",
        "Назвать результат занятия: сегодня объект начнет сам защищать свои данные от неправильных изменений.",
    ],
)

doc.add_heading("2. Теоретическая часть (25 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя:")
doc.add_paragraph(
    "Объяснять тему через одну идею: важные данные не должны свободно ломаться снаружи. Сначала показать знакомое `public`-поле, "
    "затем ввести `private` как внутреннюю часть объекта. После этого показать, что доступ все равно нужен, но через правила: "
    "методы и свойства."
)

doc.add_heading("public и private", level=4)
doc.add_paragraph(
    "`public` означает, что к полю или методу можно обратиться снаружи объекта. Это удобно, но опасно для важных данных: "
    "любой код может поставить здоровью героя отрицательное число или слишком странное значение."
)
doc.add_paragraph(
    "`private` означает, что часть класса доступна только внутри самого класса. Простая формулировка для учеников: "
    "`private` — это внутренности объекта, которые снаружи нельзя менять напрямую."
)

doc.add_heading("Инкапсуляция", level=4)
doc.add_paragraph(
    "Инкапсуляция — это принцип ООП: объект хранит свои данные внутри и сам решает, как их можно менять. "
    "Мы не лезем руками внутрь объекта, а просим объект выполнить действие через метод или свойство."
)
doc.add_paragraph(
    "Пример из игры: здоровье персонажа нельзя просто сделать `-999`. Герой должен получать урон через метод `TakeDamage()`, "
    "а внутри метода будет проверка, что здоровье не ушло ниже нуля."
)

doc.add_heading("Безопасные методы", level=4)
doc.add_paragraph(
    "Безопасный метод меняет данные по правилу. Например, `TakeDamage(int damage)` уменьшает здоровье, но затем проверяет, "
    "не стало ли оно меньше нуля. `Heal(int value)` может лечить героя только положительным значением."
)
doc.add_paragraph(
    "Такой метод похож на охранника у двери: снаружи мы просим выполнить действие, а объект внутри проверяет, можно ли его делать."
)

doc.add_heading("Свойства", level=4)
doc.add_paragraph(
    "Свойство в C# выглядит почти как поле, но может управлять чтением и изменением значения. "
    "Через `get` объект отдает значение, а через `set` может проверить новое значение перед сохранением."
)
doc.add_paragraph(
    "На этом уроке достаточно базовой идеи: поле часто хранит данные внутри класса, а свойство дает удобный и контролируемый доступ."
)

doc.add_heading("Записи в блокнот", level=4)
for term, definition in [
    ("`public`", "доступно снаружи класса."),
    ("`private`", "доступно только внутри самого класса."),
    ("Инкапсуляция", "подход, при котором объект хранит данные внутри и меняет их по своим правилам."),
    ("Безопасный метод", "метод, который меняет данные объекта и проверяет, что значение осталось правильным."),
    ("Свойство", "удобный доступ к данным объекта через `get` и `set`."),
    ("`value`", "новое значение, которое пытаются записать в свойство через `set`."),
]:
    paragraph = doc.add_paragraph()
    paragraph.add_run(term).bold = True
    paragraph.add_run(" — " + definition)

doc.add_heading("3. Практическая работа (50 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя:")
doc.add_paragraph(
    "Практику вести как постепенное улучшение одного класса. Сначала преподаватель показывает ошибочный вариант с `public Health`, "
    "затем группа вместе закрывает поле через `private`, после этого добавляет методы `TakeDamage()` и `Heal()`. В финале ученики "
    "добавляют свойство `Health` с проверкой и читают ошибку Visual Studio при попытке обратиться к закрытому полю."
)

doc.add_heading("Задание", level=4)
doc.add_paragraph("Часть 1. Увидеть проблему открытого поля.")
add_numbered(
    doc,
    [
        "Открыть Visual Studio 2022.",
        "Создать новый проект `Console App` с именем `OopLesson4` или открыть учебный проект.",
        "Открыть файл `Program.cs`.",
        "Создать класс `Player` с открытыми полями `Name` и `Health`.",
        "Создать объект `player` через конструктор.",
        "Напрямую изменить `player.Health` на отрицательное значение и запустить программу.",
    ],
)
doc.add_paragraph("[СКРИНШОТ: Visual Studio 2022 — Program.cs с public Health и строкой player.Health = -999]")
doc.add_paragraph("Код для демонстрации проблемы:")
add_code(
    doc,
    '''
Player player = new Player("Рыцарь", 100);
player.Health = -999;
player.ShowInfo();

class Player
{
    public string Name;
    public int Health;

    public Player(string name, int health)
    {
        Name = name;
        Health = health;
    }

    public void ShowInfo()
    {
        Console.WriteLine("Игрок: " + Name);
        Console.WriteLine("Здоровье: " + Health);
    }
}
''',
)
doc.add_paragraph("[СКРИНШОТ: консоль — у игрока выведено отрицательное здоровье после прямого изменения поля]")

doc.add_paragraph("Часть 2. Закрыть поле и менять здоровье через методы.")
add_numbered(
    doc,
    [
        "Заменить `public int Health` на закрытое поле `private int health`.",
        "В конструкторе записывать начальное здоровье в `health`.",
        "В `ShowInfo()` выводить `health`.",
        "Добавить методы `TakeDamage(int damage)` и `Heal(int value)`.",
        "Попробовать написать `player.health = 50` снаружи класса и посмотреть ошибку Visual Studio.",
        "Убрать ошибочную строку и вызывать только методы.",
    ],
)
doc.add_paragraph("Код после закрытия поля:")
add_code(
    doc,
    '''
Player player = new Player("Рыцарь", 100);
player.TakeDamage(30);
player.Heal(10);
player.ShowInfo();

class Player
{
    public string Name;
    private int health;

    public Player(string name, int health)
    {
        Name = name;
        this.health = health;
    }

    public void TakeDamage(int damage)
    {
        health = health - damage;
    }

    public void Heal(int value)
    {
        health = health + value;
    }

    public void ShowInfo()
    {
        Console.WriteLine("Игрок: " + Name);
        Console.WriteLine("Здоровье: " + health);
    }
}
''',
)
doc.add_paragraph("[СКРИНШОТ: Visual Studio 2022 — ошибка доступа при попытке написать player.health = 50 вне класса Player]")

doc.add_paragraph("Часть 3. Добавить проверку, чтобы здоровье не уходило ниже нуля.")
add_numbered(
    doc,
    [
        "В методе `TakeDamage()` после уменьшения здоровья добавить `if (health < 0)`.",
        "Если здоровье стало меньше нуля, вернуть его к `0`.",
        "В методе `Heal()` добавить проверку: лечить только если `value > 0`.",
        "Запустить код с большим уроном и проверить, что здоровье становится `0`, а не отрицательным.",
    ],
)
doc.add_paragraph("Код с безопасными методами:")
add_code(
    doc,
    '''
Player player = new Player("Рыцарь", 100);
player.TakeDamage(150);
player.Heal(20);
player.ShowInfo();

class Player
{
    public string Name;
    private int health;

    public Player(string name, int health)
    {
        Name = name;
        this.health = health;
    }

    public void TakeDamage(int damage)
    {
        health = health - damage;

        if (health < 0)
        {
            health = 0;
        }
    }

    public void Heal(int value)
    {
        if (value > 0)
        {
            health = health + value;
        }
    }

    public void ShowInfo()
    {
        Console.WriteLine("Игрок: " + Name);
        Console.WriteLine("Здоровье: " + health);
    }
}
''',
)
doc.add_paragraph("[СКРИНШОТ: консоль — после большого урона здоровье не уходит ниже 0]")

doc.add_paragraph("Часть 4. Добавить свойства для удобного доступа.")
add_numbered(
    doc,
    [
        "Заменить открытое поле имени на свойство `public string Name { get; private set; }`.",
        "Добавить свойство `Health` с `get` и закрытым `private set`.",
        "В `private set` проверять значение перед сохранением в поле `health`.",
        "В конструкторе и методах использовать свойство `Health`, чтобы проверка работала в одном месте.",
        "Запустить программу и сравнить: снаружи можно читать `player.Health`, но нельзя свободно записывать опасное значение.",
    ],
)
doc.add_paragraph("Итоговый код урока:")
add_code(
    doc,
    '''
Player player = new Player("Рыцарь", 100);

Console.WriteLine("Стартовое здоровье: " + player.Health);
player.TakeDamage(150);
Console.WriteLine("После урона: " + player.Health);
player.Heal(20);
Console.WriteLine("После лечения: " + player.Health);

class Player
{
    private int health;

    public string Name { get; private set; }

    public int Health
    {
        get
        {
            return health;
        }
        private set
        {
            if (value < 0)
            {
                health = 0;
            }
            else
            {
                health = value;
            }
        }
    }

    public Player(string name, int health)
    {
        Name = name;
        Health = health;
    }

    public void TakeDamage(int damage)
    {
        Health = Health - damage;
    }

    public void Heal(int value)
    {
        if (value > 0)
        {
            Health = Health + value;
        }
    }
}
''',
)
doc.add_paragraph("[СКРИНШОТ: Visual Studio 2022 — итоговый класс Player с private health, свойствами Name и Health, методами TakeDamage() и Heal()]")
add_label_paragraph(
    doc,
    "Ожидаемый результат: ",
    "ученик видит проблему прямого доступа к полю, закрывает важные данные через `private`, меняет здоровье через методы, "
    "добавляет свойство `Health` с проверкой и может объяснить, почему объект сам контролирует свое состояние.",
)

doc.add_heading("4. Самостоятельная работа (30 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя:")
doc.add_paragraph(
    "Первые 5 минут дать ученикам работать по аналогии без готового решения. На 10-й минуте проверить три вещи: есть ли `private`-поле, "
    "есть ли метод безопасного изменения, есть ли хотя бы одна проверка через `if`. Если большая часть группы застряла, вернуть их к примеру `Player` "
    "и попросить заменить только названия класса и полей."
)

doc.add_heading("Задание", level=4)
doc.add_paragraph("Выбери один вариант и сделай класс с защищенными данными.")
add_bullets(
    doc,
    [
        "Вариант 1: `Enemy` — закрытое поле `health`, методы `TakeDamage()` и `ShowInfo()`.",
        "Вариант 2: `Item` — закрытое поле `price`, свойство `Price` с проверкой, что цена не меньше 0.",
        "Вариант 3: `Robot` — закрытое поле `battery`, методы `UseEnergy()` и `Charge()`.",
        "Вариант 4: `BankAccount` — закрытое поле `balance`, методы `Deposit()` и `Spend()`.",
    ],
)
doc.add_paragraph("Минимальные требования:")
add_numbered(
    doc,
    [
        "Создать класс и конструктор.",
        "Добавить минимум одно `private`-поле.",
        "Добавить минимум одно свойство или безопасный метод для работы с этим полем.",
        "Добавить проверку значения через `if`.",
        "Создать объект класса и проверить работу в консоли.",
        "Показать преподавателю, где данные закрыты, а где разрешенный доступ к ним.",
    ],
)
doc.add_paragraph("Каркас для варианта `Item`:")
add_code(
    doc,
    '''
Item sword = new Item("Меч", 120);
Console.WriteLine(sword.Name + ": " + sword.Price);

class Item
{
    private int price;

    public string Name { get; private set; }

    public int Price
    {
        get
        {
            return price;
        }
        set
        {
            if (value >= 0)
            {
                price = value;
            }
        }
    }

    public Item(string name, int price)
    {
        Name = name;
        Price = price;
    }
}
''',
)

doc.add_heading("Критерии оценки", level=4)
add_table(
    doc,
    ["Результат", "Оценка"],
    [
        ["Есть класс, конструктор, private-поле, свойство или безопасный метод, проверка значения; программа запускается", "Отлично"],
        ["Основная защита данных сделана, но есть небольшие ошибки в названиях, выводе или расположении проверки", "Хорошо"],
        ["private-поле или проверка сделаны частично; программа заработала после помощи преподавателя", "Удовлетворительно"],
        ["Программа не запускается или данные по-прежнему меняются только напрямую через public-поле", "Требует доработки"],
    ],
)

doc.add_heading("5. Подведение итогов (10 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя:")
add_bullets(
    doc,
    [
        "Попросить нескольких учеников показать строку с `private`-полем и объяснить, почему оно закрыто.",
        "Сравнить две записи: `player.Health = -999` и `player.TakeDamage(30)`. Спросить, какая безопаснее и почему.",
        "Закрепить: объект не просто хранит данные, он управляет правилами изменения этих данных.",
        "Показать, что свойства в C# дают удобный доступ через точку, но внутри могут содержать проверку.",
        "Анонсировать следующий урок: когда объектов становится много, будем хранить их в `List<T>` и проходить по ним через `foreach`.",
    ],
)
add_label_paragraph(doc, "Вопросы для рефлексии:")
add_bullets(
    doc,
    [
        "Чем `public` отличается от `private`?",
        "Почему здоровье лучше не хранить в открытом поле?",
        "Как метод `TakeDamage()` защищает объект от неправильного состояния?",
        "Для чего в свойстве нужны `get` и `set`?",
        "Что означает `value` внутри `set`?",
    ],
)

doc.add_heading("Домашнее задание", level=2)
doc.add_paragraph("Вариант без компьютера:")
add_numbered(
    doc,
    [
        "Придумать класс для игры или жизни: персонаж, предмет, счет, робот или свой вариант.",
        "Записать 3 данных объекта.",
        "Отметить, какие данные нельзя менять напрямую.",
        "Для одного закрытого поля написать правило изменения: например, здоровье не меньше 0, цена не меньше 0, заряд от 0 до 100.",
        "Записать имя метода или свойства, через которое можно безопасно менять это значение.",
    ],
)
doc.add_paragraph("Вариант с компьютером:")
add_numbered(
    doc,
    [
        "Открыть проект `OopLesson4`.",
        "Создать свой класс с одним `private`-полем.",
        "Добавить безопасный метод или свойство с проверкой.",
        "Создать объект и проверить в консоли, что неправильное значение не ломает объект.",
    ],
)

doc.add_heading("Методические заметки преподавателя", level=2)
doc.add_heading("Возможные сложности", level=3)
add_bullets(
    doc,
    [
        "Ученики могут воспринимать `private` как запрет вообще использовать данные. Подчеркивайте: данные не исчезли, объект просто защищает доступ к ним.",
        "После замены `public Health` на `private health` ученики будут пытаться писать `player.health`. Покажите ошибку доступа как подсказку компилятора.",
        "Ученики могут забыть заменить `Health` на `health` внутри класса или наоборот. Помогает правило: закрытое поле часто пишем с маленькой буквы, свойство с большой.",
        "В свойствах ученики часто путаются в фигурных скобках `get` и `set`. Не требуйте быстрый набор с нуля, разрешайте сверяться с образцом.",
        "Слово `value` может выглядеть как новая магия. Достаточно объяснения: это новое значение, которое пытаются записать в свойство.",
        "Проверка может стоять не там: например, ученик проверяет значение после вывода в консоль, а не перед сохранением в поле.",
        "Слабые навыки Windows проявятся в том, что ученик откроет не тот проект или отдельный `Program.cs`. Проверяйте путь через `Solution Explorer` и наличие файла `.sln` при открытии проекта.",
    ],
)

doc.add_heading("Способы помощи", level=3)
add_bullets(
    doc,
    [
        "Если ученик не понимает `private`, спросите: какое значение нельзя позволять менять всем подряд?",
        "Если ученик застрял на методе, дайте первый уровень подсказки: метод должен получать число в параметре и менять закрытое поле внутри класса.",
        "Если не получается проверка, спросите: что должно случиться, если здоровье стало меньше нуля?",
        "Если ошибка доступа пугает ученика, попросите прочитать ее как сообщение: снаружи нельзя обратиться к закрытому полю.",
        "Если ученик путает поле и свойство, попросите обвести поле, где реально хранится число, и свойство, через которое число читают или меняют.",
        "Если код не запускается, сначала проверьте пары фигурных скобок, потом первую ошибку в списке Visual Studio, а не все ошибки сразу.",
        "Если ученик просит готовый код, дайте только каркас класса с пустыми телами методов, чтобы он сам дописал проверку.",
    ],
)

doc.add_heading("Дополнительные задания для тех, кто справился раньше", level=3)
add_bullets(
    doc,
    [
        "Добавить ограничение: здоровье не может быть больше 100.",
        "Добавить свойство `IsAlive`, которое возвращает `true`, если здоровье больше 0.",
        "Добавить метод `UseEnergy(int value)` для класса `Robot`, чтобы заряд не уходил ниже 0.",
        "Добавить свойство `Price` для предмета и проверить, что отрицательная цена не сохраняется.",
        "Подготовить короткое объяснение для соседа: почему `private`-поле и public-метод вместе лучше, чем просто public-поле.",
    ],
)

doc.save(OUTPUT)
print(OUTPUT)
