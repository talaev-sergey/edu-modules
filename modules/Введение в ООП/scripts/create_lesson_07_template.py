from docx import Document
from docx.shared import Pt


OUTPUT = "content/lesson_07.docx"


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    return table


def add_label_paragraph(document, label, text=""):
    paragraph = document.add_paragraph()
    paragraph.add_run(label).bold = True
    if text:
        paragraph.add_run(text)
    return paragraph


def add_code(document, code):
    for line in code.strip("\n").splitlines():
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10)


def add_numbered(document, items):
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


doc = Document()
doc.add_heading("Урок 7 — Мини-проект: простая текстовая RPG-сцена", level=1)

doc.add_heading("Общая информация", level=2)
add_table(
    doc,
    ["Параметр", "Значение"],
    [
        ["Курс", "Программирование на C#"],
        ["Модуль", "Введение в ООП"],
        ["Номер урока", "7"],
        ["Возраст учащихся", "12-15 лет"],
        ["Продолжительность", "120 мин"],
    ],
)

doc.add_heading("Цель урока", level=2)
doc.add_paragraph(
    "К концу урока ученики смогут спланировать и собрать консольную RPG-сцену из классов `Hero`, `Enemy` и `Item`: "
    "создать объекты через конструкторы, связать их методами `Attack()`, `TakeDamage()`, `AddItem()` и вывести завершенный сценарий в консоль."
)

doc.add_heading("План урока", level=2)
add_table(
    doc,
    ["Этап", "Время"],
    [
        ["1. Организационный момент", "5 мин"],
        ["2. Теоретическая часть", "10 мин"],
        ["3. Практическая работа", "65 мин"],
        ["4. Самостоятельная работа", "30 мин"],
        ["5. Подведение итогов", "10 мин"],
        ["Итого", "120 мин"],
    ],
)

doc.add_heading("Ход занятия", level=2)

doc.add_heading("1. Организационный момент (5 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя:")
add_bullets(
    doc,
    [
        "Проверить, что у учеников открыт Visual Studio 2022 и доступен проект прошлого урока или новый проект `OopLesson7`.",
        "Попросить открыть `Program.cs` через `Solution Explorer`. Если окна нет, открыть его через `View` -> `Solution Explorer`.",
        "Коротко повторить: классы описывают объекты, конструкторы создают объекты с данными, методы выполняют действия, `private` защищает внутреннее состояние.",
        "Показать цель урока одной фразой: сегодня отдельные идеи ООП соберутся в маленькую работающую сцену.",
        "Предупредить: это не полноценная игра, а учебный мини-проект, где важнее правильные связи между объектами, чем сложность боя.",
    ],
)

doc.add_heading("2. Теоретическая часть (10 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя:")
doc.add_paragraph(
    "Это урок в режиме B: короткая теория только задает карту проекта, а основное понимание появляется во время сборки кода. "
    "На доске или экране показать три карточки: `Hero`, `Enemy`, `Item`. Под каждой карточкой подписать данные и действия."
)

doc.add_heading("Проектирование классов", level=4)
doc.add_paragraph(
    "Перед кодом нужно решить, за что отвечает каждый класс. `Hero` хранит данные героя, умеет атаковать и получать предметы. "
    "`Enemy` хранит свое здоровье и сам обрабатывает получение урона. `Item` хранит название и описание награды."
)
doc.add_paragraph(
    "Хороший вопрос для учеников: кто должен уменьшать здоровье противника - `Program.cs` напрямую или сам объект `Enemy` через метод `TakeDamage()`? "
    "На этом уроке важно закрепить ответ: объект сам отвечает за свои данные."
)

doc.add_heading("Сценарий программы", level=4)
doc.add_paragraph(
    "Сценарий программы - это последовательность событий в `Program.cs`. Мы создаем героя, создаем противника, создаем награду, "
    "показываем вступление, герой атакует, противник теряет здоровье, после победы герой получает предмет."
)

doc.add_heading("Записи в блокнот", level=4)
for term, definition in [
    ("Проектирование классов", "планирование, какие классы нужны и за что отвечает каждый класс."),
    ("Связь объектов", "ситуация, когда один объект вызывает метод другого объекта."),
    ("Сценарий программы", "последовательность действий в `Program.cs`, которая собирает объекты в историю."),
    ("`TakeDamage()`", "метод, через который объект безопасно уменьшает свое здоровье."),
    ("Инвентарь", "список предметов, которые хранит герой."),
    ("Завершенная сцена", "маленькая программа, у которой есть начало, действие и понятный итог в консоли."),
]:
    paragraph = doc.add_paragraph()
    paragraph.add_run(term).bold = True
    paragraph.add_run(" — " + definition)

doc.add_heading("3. Практическая работа (65 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя:")
doc.add_paragraph(
    "Вести практику как сборку проекта по частям. Сначала вместе с учениками спланировать классы, затем создать `Item`, потом `Enemy`, потом `Hero`, "
    "и только после этого написать сценарий в `Program.cs`. Каждый блок должен давать видимый мини-результат, чтобы ученики не теряли нить проекта."
)

doc.add_heading("Задание", level=4)
doc.add_paragraph("Мини-теория 1. Планируем мини-проект через объекты.")
doc.add_paragraph(
    "До написания кода ученики должны увидеть структуру: `Hero` не должен хранить описание предмета, `Item` не должен атаковать, `Enemy` не должен выдавать награду сам по себе. "
    "Каждый класс отвечает за свою часть сцены."
)
add_numbered(
    doc,
    [
        "Открыть Visual Studio 2022.",
        "Создать новый проект `Console App` с именем `OopLesson7` или открыть учебный проект.",
        "Открыть файл `Program.cs`.",
        "Записать в комментариях план классов: `Hero`, `Enemy`, `Item`.",
        "Для каждого класса записать 2-3 данных и 1-2 действия.",
    ],
)
doc.add_paragraph("[СКРИНШОТ: Visual Studio 2022 — Program.cs с комментариями-планом классов Hero, Enemy и Item]")
doc.add_paragraph("План мини-проекта:")
add_code(
    doc,
    '''
// План мини-проекта:
// Hero: имя, здоровье, урон, инвентарь, атаковать противника, получить предмет.
// Enemy: имя, здоровье, урон, получить урон, проверить поражение, показать состояние.
// Item: название, описание, показать информацию.
''',
)

doc.add_paragraph("Мини-теория 2. Создаем `Item` и `Enemy`.")
doc.add_paragraph(
    "`Item` самый простой класс: он хранит название и описание. `Enemy` важнее: здоровье противника нельзя менять напрямую из `Program.cs`, "
    "поэтому для урона нужен метод `TakeDamage()`. Внутри метода стоит проверка, чтобы здоровье не стало меньше нуля."
)
add_numbered(
    doc,
    [
        "Создать класс `Item` с `Name`, `Description`, конструктором и `ShowInfo()`.",
        "Создать класс `Enemy` с `Name`, `Health`, `Damage`, конструктором и `ShowInfo()`.",
        "Добавить в `Enemy` метод `TakeDamage(int damage)`.",
        "Внутри `TakeDamage()` уменьшить здоровье и проверить, что оно не ушло ниже 0.",
        "Добавить в `Enemy` метод `IsDefeated()`, который возвращает `true`, если здоровье равно 0.",
    ],
)
doc.add_paragraph("[СКРИНШОТ: Visual Studio 2022 — классы Item и Enemy, выделены TakeDamage() и проверка Health < 0]")
doc.add_paragraph("Код классов `Item` и `Enemy`:")
add_code(
    doc,
    '''
class Item
{
    public string Name { get; private set; }
    public string Description { get; private set; }

    public Item(string name, string description)
    {
        Name = name;
        Description = description;
    }

    public void ShowInfo()
    {
        Console.WriteLine("Предмет: " + Name);
        Console.WriteLine(Description);
    }
}

class Enemy
{
    public string Name { get; private set; }
    public int Health { get; private set; }
    public int Damage { get; private set; }

    public Enemy(string name, int health, int damage)
    {
        Name = name;
        Health = health;
        Damage = damage;
    }

    public void TakeDamage(int damage)
    {
        Health = Health - damage;

        if (Health < 0)
        {
            Health = 0;
        }

        Console.WriteLine(Name + " получает " + damage + " урона.");
    }

    public bool IsDefeated()
    {
        return Health == 0;
    }

    public void ShowInfo()
    {
        Console.WriteLine(Name + ": здоровье " + Health + ", урон " + Damage);
    }
}
''',
)

doc.add_paragraph("Мини-теория 3. Создаем `Hero` с инвентарем.")
doc.add_paragraph(
    "`Hero` хранит список предметов внутри себя: `private List<Item> _items = new List<Item>();`. Это соединяет урок про списки с мини-проектом. "
    "Снаружи нельзя напрямую менять список, но можно вызвать метод `AddItem()`."
)
add_numbered(
    doc,
    [
        "Добавить в начало файла строку `using System.Collections.Generic;`.",
        "Создать класс `Hero` с `Name`, `Health`, `Damage`.",
        "Добавить приватный список предметов `_items`.",
        "Добавить метод `Attack(Enemy enemy)`.",
        "Внутри `Attack()` вывести текст атаки и вызвать `enemy.TakeDamage(Damage)`.",
        "Добавить метод `AddItem(Item item)`.",
        "Добавить метод `ShowInventory()` с `foreach` по списку предметов.",
    ],
)
doc.add_paragraph("[СКРИНШОТ: Visual Studio 2022 — класс Hero, выделены private List<Item> _items, Attack(Enemy enemy), AddItem() и ShowInventory()]")
doc.add_paragraph("Код класса `Hero`:")
add_code(
    doc,
    '''
class Hero
{
    public string Name { get; private set; }
    public int Health { get; private set; }
    public int Damage { get; private set; }

    private List<Item> _items = new List<Item>();

    public Hero(string name, int health, int damage)
    {
        Name = name;
        Health = health;
        Damage = damage;
    }

    public void Attack(Enemy enemy)
    {
        Console.WriteLine(Name + " атакует " + enemy.Name + ".");
        enemy.TakeDamage(Damage);
    }

    public void AddItem(Item item)
    {
        _items.Add(item);
        Console.WriteLine(Name + " получает предмет: " + item.Name);
    }

    public void ShowInventory()
    {
        Console.WriteLine("Инвентарь героя " + Name + ":");

        foreach (Item item in _items)
        {
            Console.WriteLine("- " + item.Name);
        }
    }
}
''',
)

doc.add_paragraph("Мини-теория 4. Собираем сцену в `Program.cs`.")
doc.add_paragraph(
    "Теперь классы готовы, но сцена еще не произошла. В верхней части `Program.cs` создаем объекты и пишем последовательность событий. "
    "Важно: `Program.cs` не должен сам уменьшать здоровье противника. Он только вызывает методы объектов."
)
add_numbered(
    doc,
    [
        "Создать объект `Hero hero = new Hero(\"Рыцарь\", 100, 25);`.",
        "Создать объект `Enemy enemy = new Enemy(\"Тренировочный противник\", 50, 8);`.",
        "Создать объект `Item reward = new Item(\"Зелье лечения\", \"Восстанавливает здоровье героя.\");`.",
        "Вывести вступительный текст сцены.",
        "Показать состояние противника через `enemy.ShowInfo()`.",
        "Два раза вызвать `hero.Attack(enemy)` и после каждого удара показать состояние противника.",
        "Через `if (enemy.IsDefeated())` выдать награду герою.",
        "Показать инвентарь через `hero.ShowInventory()`.",
    ],
)
doc.add_paragraph("[СКРИНШОТ: Visual Studio 2022 — верхняя часть Program.cs со сценарием: создание hero, enemy, reward и вызовы методов]")
doc.add_paragraph("Финальный код мини-проекта:")
add_code(
    doc,
    '''
using System.Collections.Generic;

Hero hero = new Hero("Рыцарь", 100, 25);
Enemy enemy = new Enemy("Тренировочный противник", 50, 8);
Item reward = new Item("Зелье лечения", "Восстанавливает здоровье героя.");

Console.WriteLine("Герой входит в тренировочную сцену.");
Console.WriteLine("На пути появляется противник.");
Console.WriteLine();

enemy.ShowInfo();
Console.WriteLine();

hero.Attack(enemy);
enemy.ShowInfo();
Console.WriteLine();

hero.Attack(enemy);
enemy.ShowInfo();
Console.WriteLine();

if (enemy.IsDefeated())
{
    Console.WriteLine("Противник побежден.");
    hero.AddItem(reward);
}

Console.WriteLine();
hero.ShowInventory();

class Item
{
    public string Name { get; private set; }
    public string Description { get; private set; }

    public Item(string name, string description)
    {
        Name = name;
        Description = description;
    }

    public void ShowInfo()
    {
        Console.WriteLine("Предмет: " + Name);
        Console.WriteLine(Description);
    }
}

class Enemy
{
    public string Name { get; private set; }
    public int Health { get; private set; }
    public int Damage { get; private set; }

    public Enemy(string name, int health, int damage)
    {
        Name = name;
        Health = health;
        Damage = damage;
    }

    public void TakeDamage(int damage)
    {
        Health = Health - damage;

        if (Health < 0)
        {
            Health = 0;
        }

        Console.WriteLine(Name + " получает " + damage + " урона.");
    }

    public bool IsDefeated()
    {
        return Health == 0;
    }

    public void ShowInfo()
    {
        Console.WriteLine(Name + ": здоровье " + Health + ", урон " + Damage);
    }
}

class Hero
{
    public string Name { get; private set; }
    public int Health { get; private set; }
    public int Damage { get; private set; }

    private List<Item> _items = new List<Item>();

    public Hero(string name, int health, int damage)
    {
        Name = name;
        Health = health;
        Damage = damage;
    }

    public void Attack(Enemy enemy)
    {
        Console.WriteLine(Name + " атакует " + enemy.Name + ".");
        enemy.TakeDamage(Damage);
    }

    public void AddItem(Item item)
    {
        _items.Add(item);
        Console.WriteLine(Name + " получает предмет: " + item.Name);
    }

    public void ShowInventory()
    {
        Console.WriteLine("Инвентарь героя " + Name + ":");

        foreach (Item item in _items)
        {
            Console.WriteLine("- " + item.Name);
        }
    }
}
''',
)
doc.add_paragraph("[СКРИНШОТ: консоль — завершенная сцена: вступление, атаки, здоровье противника, победа и инвентарь]")
add_label_paragraph(
    doc,
    "Ожидаемый результат:",
    " у ученика есть рабочая консольная сцена из классов `Hero`, `Enemy` и `Item`, где объекты связаны через методы, а не через прямое изменение чужих данных.",
)

doc.add_heading("4. Самостоятельная работа (30 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя:")
doc.add_paragraph(
    "Дать ученикам доработать сцену по выбору. Первые 5 минут не подсказывать готовый код: пусть ученик выберет, какую часть меняет. "
    "Затем проверить, что изменение не ломает главную идею: данные меняются через методы объектов, а не напрямую из `Program.cs`."
)

doc.add_heading("Задание", level=4)
doc.add_paragraph("Выбери одно улучшение мини-проекта и реализуй его.")
add_bullets(
    doc,
    [
        "Вариант 1: изменить героя, противника и награду, чтобы сцена получила свой сюжет.",
        "Вариант 2: добавить второй предмет-награду и положить его в инвентарь через `AddItem()`.",
        "Вариант 3: добавить метод `Heal(int value)` в `Hero` и вызвать его после получения предмета.",
        "Вариант 4: добавить метод `AttackHero(Hero hero)` в `Enemy` как подготовку к следующему уроку.",
        "Вариант 5: добавить второй вызов `reward.ShowInfo()` после получения предмета.",
    ],
)
doc.add_paragraph("Минимальные требования:")
add_numbered(
    doc,
    [
        "Проект запускается.",
        "В проекте есть классы `Hero`, `Enemy`, `Item`.",
        "Объекты создаются через конструкторы.",
        "Урон противнику проходит через `TakeDamage()`.",
        "Предмет добавляется через `AddItem()`.",
        "В консоли видно начало, действие и итог сцены.",
    ],
)
doc.add_paragraph("Каркас для варианта с лечением героя:")
add_code(
    doc,
    '''
public void Heal(int value)
{
    if (value > 0)
    {
        Health = Health + value;
        Console.WriteLine(Name + " восстанавливает " + value + " здоровья.");
    }
}
''',
)

doc.add_heading("Критерии оценки", level=4)
add_table(
    doc,
    ["Результат", "Оценка"],
    [
        ["Собрана рабочая сцена с `Hero`, `Enemy`, `Item`; есть конструкторы, методы `Attack()`, `TakeDamage()`, `AddItem()`, инвентарь и понятный вывод в консоль", "Отлично"],
        ["Основная сцена работает, но есть небольшие ошибки в тексте вывода, названиях или оформлении кода", "Хорошо"],
        ["Классы или сценарий сделаны частично; проект заработал после подсказок преподавателя", "Удовлетворительно"],
        ["Проект не запускается или объекты почти не связаны методами", "Требует доработки"],
    ],
)

doc.add_heading("5. Подведение итогов (10 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя:")
add_bullets(
    doc,
    [
        "Попросить 2-3 учеников показать свою сцену и объяснить, за что отвечает каждый класс.",
        "Вернуться к ключевой идее: `Program.cs` управляет сценарием, но сами действия живут внутри классов.",
        "Спросить, почему `Enemy` сам уменьшает свое здоровье через `TakeDamage()`.",
        "Спросить, почему инвентарь хранится внутри `Hero`, а добавление идет через `AddItem()`.",
        "Анонсировать следующий урок: будем дорабатывать проект, добавим простой выбор действия и несколько противников в списке.",
    ],
)
add_label_paragraph(doc, "Вопросы для рефлексии:")
add_bullets(
    doc,
    [
        "Какие три основных класса есть в мини-проекте?",
        "Какая ответственность у `Hero`?",
        "Какая ответственность у `Enemy`?",
        "Какая ответственность у `Item`?",
        "Почему лучше вызвать `enemy.TakeDamage(Damage)`, чем менять `enemy.Health` напрямую?",
        "Где в проекте используется список?",
    ],
)

doc.add_heading("Домашнее задание", level=2)
doc.add_paragraph("Вариант без компьютера:")
add_numbered(
    doc,
    [
        "Нарисовать схему классов `Hero`, `Enemy`, `Item`.",
        "Под каждым классом записать его данные и действия.",
        "Придумать свой сценарий сцены из 5-7 событий.",
        "Отметить, где герой атакует противника, где противник получает урон, где герой получает предмет.",
        "Придумать одно улучшение для итогового урока: выбор действия, несколько противников или новый предмет.",
    ],
)
doc.add_paragraph("Вариант с компьютером:")
add_numbered(
    doc,
    [
        "Открыть проект `OopLesson7`.",
        "Доработать текст сцены, названия объектов и награду.",
        "Добавить минимум один новый метод или один новый предмет.",
        "Проверить, что проект запускается без ошибок.",
        "Подготовить короткое объяснение: какие классы есть в проекте и как они связаны.",
    ],
)

doc.add_heading("Методические заметки преподавателя", level=2)
doc.add_heading("Возможные сложности", level=3)
add_bullets(
    doc,
    [
        "Ученики могут пытаться написать всю сцену только в `Program.cs` без классов. Возвращайте к вопросу: какой объект должен отвечать за это действие?",
        "При копировании классов ученики часто теряют фигурные скобки. Если ошибок много, сначала проверить границы классов `Item`, `Enemy`, `Hero`.",
        "В `TakeDamage()` ученики могут забыть проверку `Health < 0`, и здоровье станет отрицательным.",
        "В `Hero.Attack(Enemy enemy)` ученики могут написать только текст атаки и забыть вызвать `enemy.TakeDamage(Damage)`.",
        "В инвентаре часто забывают строку `using System.Collections.Generic;` или создают список без `new List<Item>()`.",
        "Ученики могут напрямую менять `Health` из `Program.cs`. Подчеркните: на этом уроке здоровье противника меняется только через метод объекта.",
        "Слабые навыки Windows проявятся в открытии не того проекта или отдельного файла вместо решения. Проверяйте `Solution Explorer`, имя проекта и файл `Program.cs`.",
    ],
)

doc.add_heading("Способы помощи", level=3)
add_bullets(
    doc,
    [
        "Если ученик не понимает структуру проекта, попросите устно ответить: кто атакует, кто получает урон, кто хранит предмет?",
        "Если ученик застрял на `TakeDamage()`, задайте подсказку: какое число должно измениться и какая проверка нужна после изменения?",
        "Если не работает инвентарь, проверьте три вещи: `using System.Collections.Generic;`, строку `private List<Item> _items = new List<Item>();`, цикл `foreach`.",
        "Если ученик смешивает обязанности классов, предложите карточки: `Hero` делает только действия героя, `Enemy` только действия противника, `Item` только данные предмета.",
        "Если код не запускается, начать с первой ошибки Visual Studio сверху, а не читать весь список ошибок сразу.",
        "Если ученик просит готовую доработку, дайте только каркас метода и попросите самому выбрать текст вывода и место вызова.",
    ],
)

doc.add_heading("Дополнительные задания для тех, кто справился раньше", level=3)
add_bullets(
    doc,
    [
        "Добавить метод `ShowStatus()` в `Hero` и вывести здоровье героя.",
        "Добавить второй предмет в инвентарь и вывести оба предмета через `ShowInventory()`.",
        "Добавить метод `UseItem(string itemName)` как заготовку без удаления предмета из списка.",
        "Добавить второго противника как отдельный объект, но пока без сложного меню.",
        "Разнести классы по отдельным файлам через `Project` -> `Add Class`, если группа уверенно работает с файлами Visual Studio.",
    ],
)

doc.save(OUTPUT)
print(OUTPUT)
