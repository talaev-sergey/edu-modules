from pathlib import Path

from docx import Document


SOURCE_DIR = Path(r"D:\EduDocs\Введение в ООП")
OUTPUT = Path("content/source_oop_10_lessons.md")

FILES = [
    "Учебная программа ООП на C#.docx",
    "vvedenie_v_oop_zanyatie_1.docx",
    "vvedenie_v_oop_zanyatie_2.docx",
    "vvedenie_v_oop_zanyatie_3.docx",
    "vvedenie_v_oop_zanyatie_4.docx",
    "vvedenie_v_oop_zanyatie_5.docx",
    "vvedenie_v_oop_zanyatie_6.docx",
    "vvedenie_v_oop_zanyatie_7.docx",
    "vvedenie_v_oop_zanyatie_8.docx",
    "vvedenie_v_oop_zanyatie_9.docx",
    "vvedenie_v_oop_zanyatie_10.docx",
]


def table_to_lines(table):
    lines = []
    for row in table.rows:
        values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        if any(values):
            lines.append("| " + " | ".join(values) + " |")
    return lines


def extract_docx(path):
    document = Document(path)
    lines = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for index, table in enumerate(document.tables, start=1):
        lines.append(f"[Таблица {index}]")
        lines.extend(table_to_lines(table))
    return lines


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    result = ["# Исходные материалы модуля ООП на C#", ""]
    for file_name in FILES:
        path = SOURCE_DIR / file_name
        result.append(f"## {file_name}")
        result.append("")
        if not path.exists():
            result.append("Файл не найден.")
            result.append("")
            continue
        lines = extract_docx(path)
        result.extend(lines)
        result.append("")
    OUTPUT.write_text("\n".join(result), encoding="utf-8")
    print(OUTPUT)


if __name__ == "__main__":
    main()
