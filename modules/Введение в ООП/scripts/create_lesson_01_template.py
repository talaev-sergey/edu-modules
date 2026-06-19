from docx import Document


OUTPUT = "content/lesson_01.docx"
DASH = "–"


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    return table


def add_label_paragraph(document, label, text=DASH):
    paragraph = document.add_paragraph()
    paragraph.add_run(label).bold = True
    paragraph.add_run(text)
    return paragraph


doc = Document()
doc.add_heading("Урок 1 — Классы и объекты: создаем первую модель", level=1)

doc.add_heading("Общая информация", level=2)
add_table(
    doc,
    ["Параметр", "Значение"],
    [
        ["Курс", "Программирование на C#"],
        ["Модуль", "Введение в ООП"],
        ["Номер урока", "1"],
        ["Возраст учащихся", "12-15 лет"],
        ["Продолжительность", "120 мин"],
    ],
)

doc.add_heading("Цель урока", level=2)
doc.add_paragraph(
    "К концу урока ученики смогут создать класс C#, создать 2-3 объекта этого класса "
    "через new и вывести значения их данных в консоль."
)

doc.add_heading("План урока", level=2)
add_table(
    doc,
    ["Этап", "Время"],
    [
        ["1. Организационный момент", "5 мин"],
        ["2. Теоретическая часть", "25 мин"],
        ["3. Практическая работа", "50 мин"],
        ["4. Самостоятельная работа", "30 мин"],
        ["5. Подведение итогов", "10 мин"],
        ["Итого", "120 мин"],
    ],
)

doc.add_heading("Ход занятия", level=2)

doc.add_heading("1. Организационный момент (5 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя: ")

doc.add_heading("2. Теоретическая часть (25 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя: ")
doc.add_heading("Класс", level=4)
doc.add_paragraph(DASH)
doc.add_heading("Объект", level=4)
doc.add_paragraph(DASH)
doc.add_heading("Поле", level=4)
doc.add_paragraph(DASH)
doc.add_heading("Записи в блокнот", level=4)
doc.add_paragraph(DASH)

doc.add_heading("3. Практическая работа (50 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя: ")
doc.add_heading("Задание", level=4)
doc.add_paragraph(DASH)
add_label_paragraph(doc, "Ожидаемый результат: ")

doc.add_heading("4. Самостоятельная работа (30 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя: ")
doc.add_heading("Задание", level=4)
doc.add_paragraph(DASH)
doc.add_heading("Критерии оценки", level=4)
add_table(
    doc,
    ["Результат", "Оценка"],
    [
        ["Выполнено полностью самостоятельно", "Отлично"],
        ["Выполнено с небольшими ошибками или подсказкой", "Хорошо"],
        ["Выполнено частично, потребовалась помощь", "Удовлетворительно"],
        ["Не выполнено", "Требует доработки"],
    ],
)

doc.add_heading("5. Подведение итогов (10 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя: ")
doc.add_heading("Вопросы для рефлексии", level=4)
for question in [
    "Что нового узнали сегодня?",
    "Что было самым сложным?",
    "Где это можно применить в жизни?",
]:
    doc.add_paragraph(question, style="List Bullet")

doc.add_heading("Домашнее задание", level=2)
doc.add_paragraph(DASH)

doc.add_heading("Методические заметки преподавателя", level=2)
doc.add_heading("Возможные сложности", level=3)
doc.add_paragraph(DASH)
doc.add_heading("Способы помощи", level=3)
doc.add_paragraph(DASH)
doc.add_heading("Дополнительные задания для тех, кто справился раньше", level=3)
doc.add_paragraph(DASH)

doc.save(OUTPUT)
print(OUTPUT)
