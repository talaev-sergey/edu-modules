from docx import Document
from docx.shared import Pt


OUTPUT = "content/lesson_05.docx"


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    return table


def add_label_paragraph(document, label, text=""):
    paragraph = document.add_paragraph()
    paragraph.add_run(label).bold = True
    if text:
        paragraph.add_run(text)
    return paragraph


def add_code(document, code):
    for line in code.strip("\n").splitlines():
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10)


def add_numbered(document, items):
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


doc = Document()
doc.add_heading("Урок 5 — Списки объектов: когда объектов становится много", level=1)

doc.add_heading("Общая информация", level=2)
add_table(
    doc,
    ["Параметр", "Значение"],
    [
        ["Курс", "Программирование на C#"],
        ["Модуль", "Введение в ООП"],
        ["Номер урока", "5"],
        ["Возраст учащихся", "12-15 лет"],
        ["Продолжительность", "120 мин"],
    ],
)

doc.add_heading("Цель урока", level=2)
doc.add_paragraph(
    "К концу урока ученики смогут создать `List<Hero>`, добавить в него минимум три объекта через `Add()`, "
    "пройти по списку через `foreach` и найти один объект по простому условию."
)

doc.add_heading("План урока", level=2)
add_table(
    doc,
    ["Этап", "Время"],
    [
        ["1. Организационный момент", "5 мин"],
        ["2. Теоретическая часть", "25 мин"],
        ["3. Практическая работа", "50 мин"],
        ["4. Самостоятельная работа", "30 мин"],
        ["5. Подведение итогов", "10 мин"],
        ["Итого", "120 мин"],
    ],
)

doc.add_heading("Ход занятия", level=2)

doc.add_heading("1. Организационный момент (5 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя:")
add_bullets(
    doc,
    [
        "Проверить, что у учеников открыт Visual Studio 2022 и доступен учебный проект прошлого урока или новый проект `OopLesson5`.",
        "Попросить открыть `Program.cs` через `Solution Explorer`. Если окна нет, открыть его через `View` -> `Solution Explorer`.",
        "Коротко повторить: класс `Hero` описывает общий тип объекта, конструктор задает начальные данные, свойства могут защищать данные от неправильного изменения.",
        "Задать вопрос урока: что будет, если героев, предметов или врагов станет не один и не два, а десять или двадцать?",
        "Назвать результат занятия: сегодня ученики научатся складывать объекты одного типа в список и выполнять действие для каждого объекта по очереди.",
    ],
)

doc.add_heading("2. Теоретическая часть (25 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя:")
doc.add_paragraph(
    "Вести объяснение от знакомого к новому. Сначала показать три отдельные переменные `knight`, `mage`, `archer` и повторяющиеся вызовы `ShowInfo()`. "
    "Затем объяснить, что список помогает хранить такие объекты вместе. Не вводить сложные термины про коллекции и обобщения: для этого урока достаточно идеи "
    "\"список героев\" и практики с `Add()` и `foreach`."
)

doc.add_heading("Зачем нужен список объектов", level=4)
doc.add_paragraph(
    "Когда объектов мало, их можно хранить в отдельных переменных: `knight`, `mage`, `archer`. Но если в игре появляется команда, инвентарь или набор врагов, "
    "ручной вызов метода для каждого объекта быстро становится неудобным. Список решает эту проблему: объекты лежат в одном месте, а программа может пройти по ним по очереди."
)
doc.add_paragraph(
    "Пример из игры: в инвентаре может быть меч, зелье и ключ. Игрок не пишет отдельный код для каждого предмета каждый раз, а хранит предметы в списке и показывает их одним циклом."
)

doc.add_heading("List<T>", level=4)
doc.add_paragraph(
    "`List<T>` в C# означает список объектов определенного типа. Вместо `T` пишут тип объектов, которые будут лежать в списке: `List<Hero>` хранит героев, "
    "`List<Item>` хранит предметы. Простая формулировка для учеников: внутри угловых скобок написано, кого можно класть в этот список."
)
doc.add_paragraph(
    "Если Visual Studio подчеркивает `List` красным, чаще всего нужно добавить в начало файла строку `using System.Collections.Generic;` или принять подсказку Visual Studio."
)

doc.add_heading("Add()", level=4)
doc.add_paragraph(
    "`Add()` добавляет объект в список. Можно сначала создать объект в отдельной переменной и добавить его, а можно сразу написать `heroes.Add(new Hero(...))`. "
    "Для новичков удобнее начать с готового примера, а затем показать более короткую запись."
)
doc.add_paragraph(
    "Важно проговорить ограничение: в `List<Hero>` можно добавить объект `Hero`, но нельзя добавить объект `Item`, потому что это другой тип."
)

doc.add_heading("foreach", level=4)
doc.add_paragraph(
    "`foreach` проходит по списку по одному объекту за раз. В строке `foreach (Hero hero in heroes)` слово `hero` означает текущего героя, с которым цикл работает прямо сейчас. "
    "Внутри фигурных скобок можно вызвать метод у каждого объекта, например `hero.ShowInfo()`."
)
doc.add_paragraph(
    "Это удобнее, чем писать `knight.ShowInfo()`, `mage.ShowInfo()`, `archer.ShowInfo()` вручную. Если добавить в список четвертого героя, цикл сам пройдет и по нему тоже."
)

doc.add_heading("Записи в блокнот", level=4)
for term, definition in [
    ("`List<T>`", "список объектов одного типа."),
    ("`List<Hero>`", "список, в котором лежат только объекты `Hero`."),
    ("`Add()`", "метод, который добавляет объект в список."),
    ("`foreach`", "цикл, который проходит по каждому объекту списка."),
    ("`using System.Collections.Generic;`", "строка, которая подключает `List`, если Visual Studio его не видит."),
    ("Поиск по условию", "проход по списку, при котором программа выбирает объект с нужным значением."),
]:
    paragraph = doc.add_paragraph()
    paragraph.add_run(term).bold = True
    paragraph.add_run(" — " + definition)

doc.add_heading("3. Практическая работа (50 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя:")
doc.add_paragraph(
    "Практику вести по схеме \"показал - сделал вместе - отпустил часть работы ученикам\". Сначала преподаватель показывает полную версию с отдельными переменными, "
    "затем вместе с группой переносит объекты в `List<Hero>`, потом ученики дописывают новый объект и проверяют, что `foreach` выводит его без дополнительного ручного вызова. "
    "Последний блок - поиск героя с самым большим здоровьем через обычный цикл."
)

doc.add_heading("Задание", level=4)
doc.add_paragraph("Часть 1. Увидеть проблему отдельных переменных.")
add_numbered(
    doc,
    [
        "Открыть Visual Studio 2022.",
        "Создать новый проект `Console App` с именем `OopLesson5` или открыть учебный проект.",
        "Открыть файл `Program.cs`.",
        "Вставить класс `Hero` и создать трех героев: рыцаря, мага и лучника.",
        "Вызвать `ShowInfo()` для каждого героя отдельной строкой.",
        "Запустить проект кнопкой запуска в Visual Studio или сочетанием `Ctrl+F5`.",
    ],
)
doc.add_paragraph("[СКРИНШОТ: Visual Studio 2022 — Program.cs с отдельными переменными knight, mage, archer и тремя вызовами ShowInfo()]")
doc.add_paragraph("Стартовый код:")
add_code(
    doc,
    '''
Hero knight = new Hero("Рыцарь", 100, 15);
Hero mage = new Hero("Маг", 60, 30);
Hero archer = new Hero("Лучник", 80, 20);

knight.ShowInfo();
mage.ShowInfo();
archer.ShowInfo();

class Hero
{
    public string Name { get; private set; }
    public int Health { get; private set; }
    public int Damage { get; private set; }

    public Hero(string name, int health, int damage)
    {
        Name = name;
        Health = health;
        Damage = damage;
    }

    public void ShowInfo()
    {
        Console.WriteLine("Герой: " + Name);
        Console.WriteLine("Здоровье: " + Health);
        Console.WriteLine("Урон: " + Damage);
        Console.WriteLine();
    }
}
''',
)

doc.add_paragraph("Часть 2. Перенести героев в `List<Hero>`.")
add_numbered(
    doc,
    [
        "В самое начало файла добавить строку `using System.Collections.Generic;`.",
        "Удалить отдельные переменные `knight`, `mage`, `archer` или закомментировать их.",
        "Создать список `List<Hero> heroes = new List<Hero>();`.",
        "Добавить в список трех героев через `heroes.Add(new Hero(...));`.",
        "Написать `foreach (Hero hero in heroes)`.",
        "Внутри цикла вызвать `hero.ShowInfo();`.",
        "Запустить программу и сравнить вывод с первой частью.",
    ],
)
doc.add_paragraph("[СКРИНШОТ: Visual Studio 2022 — Program.cs со строкой using System.Collections.Generic, List<Hero> heroes и вызовами Add()]")
doc.add_paragraph("Код после перехода на список:")
add_code(
    doc,
    '''
using System.Collections.Generic;

List<Hero> heroes = new List<Hero>();

heroes.Add(new Hero("Рыцарь", 100, 15));
heroes.Add(new Hero("Маг", 60, 30));
heroes.Add(new Hero("Лучник", 80, 20));

foreach (Hero hero in heroes)
{
    hero.ShowInfo();
}

class Hero
{
    public string Name { get; private set; }
    public int Health { get; private set; }
    public int Damage { get; private set; }

    public Hero(string name, int health, int damage)
    {
        Name = name;
        Health = health;
        Damage = damage;
    }

    public void ShowInfo()
    {
        Console.WriteLine("Герой: " + Name);
        Console.WriteLine("Здоровье: " + Health);
        Console.WriteLine("Урон: " + Damage);
        Console.WriteLine();
    }
}
''',
)
doc.add_paragraph("[СКРИНШОТ: консоль — выведены все герои из списка через один foreach]")

doc.add_paragraph("Часть 3. Проверить удобство списка.")
add_numbered(
    doc,
    [
        "Добавить в список четвертого героя: `heroes.Add(new Hero(\"Лекарь\", 70, 10));`.",
        "Не добавлять новый ручной вызов `ShowInfo()`.",
        "Запустить программу.",
        "Убедиться, что четвертый герой тоже появился в консоли, потому что цикл проходит по всему списку.",
    ],
)

doc.add_paragraph("Часть 4. Найти героя с самым большим здоровьем.")
add_numbered(
    doc,
    [
        "После цикла вывода создать переменную `Hero strongestByHealth = heroes[0];`.",
        "Написать новый цикл `foreach (Hero hero in heroes)`.",
        "Внутри цикла сравнить `hero.Health` и `strongestByHealth.Health`.",
        "Если у текущего героя здоровья больше, записать его в `strongestByHealth`.",
        "После цикла вывести строку `Герой с самым большим здоровьем:` и вызвать `strongestByHealth.ShowInfo();`.",
    ],
)
doc.add_paragraph("[СКРИНШОТ: Visual Studio 2022 — Program.cs с циклом foreach, который обновляет strongestByHealth]")
doc.add_paragraph("Код поиска:")
add_code(
    doc,
    '''
Hero strongestByHealth = heroes[0];

foreach (Hero hero in heroes)
{
    if (hero.Health > strongestByHealth.Health)
    {
        strongestByHealth = hero;
    }
}

Console.WriteLine("Герой с самым большим здоровьем:");
strongestByHealth.ShowInfo();
''',
)
doc.add_paragraph("[СКРИНШОТ: консоль — отдельная строка Герой с самым большим здоровьем и данные найденного героя]")
add_label_paragraph(
    doc,
    "Ожидаемый результат:",
    " у ученика есть рабочий проект, где несколько объектов `Hero` хранятся в `List<Hero>`, выводятся через `foreach`, а один герой выбирается по условию.",
)

doc.add_heading("4. Самостоятельная работа (30 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя:")
doc.add_paragraph(
    "Раздать задание по выбору и сначала дать 5 минут на самостоятельный старт. Через 5 минут быстро проверить у каждого: есть ли класс, есть ли `List<...>`, "
    "есть ли хотя бы один `Add()`. Если многие ученики застряли, показать на доске каркас `Item` и попросить заменить только названия и свойства под свой вариант."
)

doc.add_heading("Задание", level=4)
doc.add_paragraph("Выбери один вариант и создай свой список объектов.")
add_bullets(
    doc,
    [
        "Вариант 1: `List<Item>` - предметы инвентаря с названием и ценой.",
        "Вариант 2: `List<Enemy>` - противники с названием, здоровьем и уроном.",
        "Вариант 3: `List<Robot>` - роботы с названием и зарядом батареи.",
        "Вариант 4: `List<Weapon>` - оружие с названием и уроном.",
    ],
)
doc.add_paragraph("Минимальные требования:")
add_numbered(
    doc,
    [
        "Создать класс с 2-3 свойствами или полями.",
        "Добавить конструктор.",
        "Создать список нужного типа: например, `List<Item> items = new List<Item>();`.",
        "Добавить минимум 3 объекта через `Add()`.",
        "Вывести все объекты через `foreach`.",
        "Проверить, что программа запускается.",
        "Дополнительное задание: найти объект по условию, например самый дорогой предмет или противника с самым большим здоровьем.",
    ],
)
doc.add_paragraph("Каркас для варианта `Item`:")
add_code(
    doc,
    '''
using System.Collections.Generic;

List<Item> items = new List<Item>();

items.Add(new Item("Меч", 120));
items.Add(new Item("Зелье", 40));
items.Add(new Item("Ключ", 10));

foreach (Item item in items)
{
    item.ShowInfo();
}

class Item
{
    public string Name { get; private set; }
    public int Price { get; private set; }

    public Item(string name, int price)
    {
        Name = name;
        Price = price;
    }

    public void ShowInfo()
    {
        Console.WriteLine("Предмет: " + Name);
        Console.WriteLine("Цена: " + Price);
        Console.WriteLine();
    }
}
''',
)

doc.add_heading("Критерии оценки", level=4)
add_table(
    doc,
    ["Результат", "Оценка"],
    [
        ["Создан список минимум из 3 объектов, объекты добавлены через Add(), вывод сделан через foreach, есть простой поиск по условию; программа запускается", "Отлично"],
        ["Список, Add() и foreach работают, но есть небольшие ошибки в названиях, выводе или оформлении кода", "Хорошо"],
        ["Класс или список сделаны частично; программа заработала после подсказок преподавателя", "Удовлетворительно"],
        ["Нет рабочего списка объектов или программа не запускается даже после базовой помощи", "Требует доработки"],
    ],
)

doc.add_heading("5. Подведение итогов (10 мин)", level=3)
add_label_paragraph(doc, "Действия преподавателя:")
add_bullets(
    doc,
    [
        "Попросить 2-3 учеников показать свой список объектов и объяснить, что написано внутри угловых скобок.",
        "Сравнить два подхода: три ручных вызова `ShowInfo()` и один цикл `foreach`.",
        "Закрепить: `Add()` добавляет объект, `foreach` проходит по каждому объекту, а `List<Hero>` хранит только героев.",
        "Показать связь с будущими уроками: на следующем занятии похожие объекты получат общего родителя через наследование.",
        "Напомнить, что на этом уроке не нужны массивы, LINQ, сортировки и `Dictionary`; главная цель - уверенно работать со списком объектов.",
    ],
)
add_label_paragraph(doc, "Вопросы для рефлексии:")
add_bullets(
    doc,
    [
        "Зачем нужен `List<Hero>`, если можно создать `knight`, `mage` и `archer` отдельно?",
        "Что означает `Hero` внутри `List<Hero>`?",
        "Что делает `Add()`?",
        "Почему `foreach` удобнее, чем несколько одинаковых вызовов `ShowInfo()`?",
        "Можно ли в `List<Hero>` положить объект `Item`? Почему?",
        "Как программа нашла героя с самым большим здоровьем?",
    ],
)

doc.add_heading("Домашнее задание", level=2)
doc.add_paragraph("Вариант без компьютера:")
add_numbered(
    doc,
    [
        "Придумать список объектов для игры или приложения: например, предметы, враги, уровни, задания, роботы или транспорт.",
        "Записать тип списка: например, `List<Item>` или `List<Robot>`.",
        "Записать минимум 3 объекта, которые будут лежать в списке.",
        "Для каждого объекта записать 2-3 данных: название, здоровье, цена, урон, заряд или свой вариант.",
        "Написать словами, что должен делать `foreach` с каждым объектом.",
    ],
)
doc.add_paragraph("Вариант с компьютером:")
add_numbered(
    doc,
    [
        "Открыть проект `OopLesson5`.",
        "Создать свой список объектов по аналогии с уроком.",
        "Добавить минимум 3 объекта через `Add()`.",
        "Вывести все объекты через `foreach`.",
        "По желанию найти объект по условию: с самой большой ценой, самым большим здоровьем или самым большим уроном.",
    ],
)

doc.add_heading("Методические заметки преподавателя", level=2)
doc.add_heading("Возможные сложности", level=3)
add_bullets(
    doc,
    [
        "Ученики могут забыть `using System.Collections.Generic;`, и `List` будет подчеркнут красным. Покажите, что Visual Studio часто предлагает исправление автоматически.",
        "В записи `List<Hero>` ученики могут забывать угловые скобки или писать `List Hero`. Сравните с коробкой, на которой написан тип содержимого.",
        "Часть учеников начнет использовать массивы или искать сложные способы. Возвращайте к цели урока: `List<T>`, `Add()` и `foreach`.",
        "При `heroes.Add(...)` ученики могут попытаться добавить объект другого типа. Проговорите: список героев принимает героев, список предметов принимает предметы.",
        "Частая ошибка - создать `List<Hero> heroes;`, но забыть `new List<Hero>()`. Объясните, что переменная списка должна получить новый пустой список.",
        "В `foreach` ученики путают `heroes` и `hero`. Помогает фраза: `heroes` - весь список, `hero` - один текущий объект из списка.",
        "В поиске по условию важно, чтобы список уже не был пустым перед строкой `Hero strongestByHealth = heroes[0];`.",
        "Слабые навыки Windows проявятся в том, что ученик откроет не тот проект или отдельный `Program.cs`. Проверяйте путь через `Solution Explorer` и наличие файла `.sln`.",
    ],
)

doc.add_heading("Способы помощи", level=3)
add_bullets(
    doc,
    [
        "Если ученик не понимает `List<Hero>`, спросите: какие объекты должны лежать в этом списке? Что написано внутри угловых скобок?",
        "Если ученик застрял на `Add()`, дайте первый уровень подсказки: список уже создан, теперь нужно положить в него новый объект.",
        "Если не получается `foreach`, спросите: как назвать один объект, который цикл достает из списка прямо сейчас?",
        "Если вывод не появляется, сначала проверьте, находится ли `ShowInfo()` внутри фигурных скобок `foreach`.",
        "Если поиск по условию не получается, начните с готовой идеи: сначала считаем первым лучшим героя из списка, потом сравниваем остальных с ним.",
        "Если ученик просит готовый код, дайте каркас с пустыми местами для типа списка, названий объектов и тела `ShowInfo()`.",
        "Если Visual Studio показывает много ошибок, разберите только первую ошибку сверху, затем снова запустите проверку.",
    ],
)

doc.add_heading("Дополнительные задания для тех, кто справился раньше", level=3)
add_bullets(
    doc,
    [
        "Добавить четвертого и пятого героя и убедиться, что `foreach` выводит всех без новых ручных вызовов.",
        "Найти героя с самым большим уроном через переменную `strongestByDamage`.",
        "Посчитать, сколько героев имеют здоровье больше 70.",
        "Создать `List<Item>` и найти самый дорогой предмет.",
        "Сделать метод `ShowTeam(List<Hero> heroes)`, который принимает список героев и выводит всю команду.",
    ],
)

doc.save(OUTPUT)
print(OUTPUT)
